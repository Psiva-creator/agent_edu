import logging
import io
from services.llm_service import LLMService
from prompts.templates import COVER_LETTER_PROMPT
from docx import Document
from xhtml2pdf import pisa

logger = logging.getLogger(__name__)

class CoverLetterAgent:
    def __init__(self):
        self.llm = LLMService()

    async def generate_cover_letter(self, data: dict) -> dict:
        """
        Generate a personalized cover letter using the LLM.
        """
        if not self.llm.is_available:
            logger.warning(
                "[Fallback] CoverLetterAgent using fallback because LLM is unavailable.",
                extra={"agent": "CoverLetterAgent", "source": "fallback", "reason": "llm_unavailable"}
            )
            return self._generate_fallback(data)

        skills = data.get("skills") or []
        if not isinstance(skills, list):
            skills = []

        prompt = COVER_LETTER_PROMPT.format(
            tone=data.get("tone", "Formal"),
            target_role=data.get("target_role", "Role"),
            resume_text=data.get("resume_text", "")[:3000],
            skills=", ".join(skills),
            experience_years=data.get("experience_years", 0),
            projects=str(data.get("projects", [])),
            job_description=data.get("job_description") or "Not provided"
        )

        try:
            result = await self.llm.generate_json(prompt)
            if isinstance(result, dict) and "paragraphs" in result:
                result["source"] = "ai"
                return result
        except Exception as e:
            error_msg = str(e).lower()
            reason = "unknown_error"
            if "quota" in error_msg or "429" in error_msg:
                reason = "quota_exhausted"
            elif "auth" in error_msg or "401" in error_msg or "403" in error_msg:
                reason = "auth_error"
            elif "timeout" in error_msg:
                reason = "timeout"
                
            logger.warning(
                f"[Fallback] CoverLetterAgent using fallback due to LLM error. Reason: {reason}. Details: {e}",
                extra={"agent": "CoverLetterAgent", "source": "fallback", "reason": reason}
            )
            
        logger.warning(
            "[Fallback] CoverLetterAgent using fallback because LLM is unavailable.",
            extra={"agent": "CoverLetterAgent", "source": "fallback", "reason": "llm_unavailable"}
        )
        return self._generate_fallback(data)
        
    def _generate_fallback(self, data: dict) -> dict:
        role = data.get("target_role", "the open position")
        tone = data.get("tone", "Formal")
        skills = data.get("skills", [])
        top_skills = ", ".join(skills[:3]) if skills else "my skills"
        exp = data.get("experience_years", 0)
        exp_text = f"{exp} years of experience" if exp > 0 else "my academic and project background"
        
        # Add dynamic tone phrasing
        if tone.lower() == "startup":
            greeting = "Hi Team,"
            closing = "I'm pumped to bring my hustle to your organization. Thanks!"
        else:
            greeting = "Dear Hiring Team,"
            closing = "I am excited about the opportunity to bring my unique blend of skills to your organization. Thank you for your time and consideration."
            
        return {
            "paragraphs": [
                {
                    "text": f"{greeting}\n\nI am writing to express my strong interest in the {role} position. With my background and passion for delivering results, I am confident in my ability to contribute effectively to your team.",
                    "explanation": "Standard opening to state intent clearly."
                },
                {
                    "text": f"Throughout my career, I have developed a robust skill set that aligns with the requirements of this role, specifically {top_skills}. With {exp_text}, my experience in leading projects and driving impact has prepared me to tackle the challenges your team is facing.",
                    "explanation": "Highlights general experience and project impact."
                },
                {
                    "text": f"{closing}",
                    "explanation": "Professional closing paragraph with a call to action."
                }
            ],
            "source": "fallback"
        }

    def generate_pdf(self, name: str, paragraphs: list[str]) -> bytes:
        """
        Convert the cover letter paragraphs to a PDF file using xhtml2pdf.
        """
        html_content = f"""
        <html>
        <head>
            <style>
                @page {{ size: a4 portrait; margin: 2cm; }}
                body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11pt; color: #333333; line-height: 1.6; }}
                h1 {{ font-size: 16pt; margin-bottom: 20px; color: #111111; border-bottom: 1px solid #eeeeee; padding-bottom: 10px; }}
                p {{ margin-bottom: 15px; }}
            </style>
        </head>
        <body>
            <h1>{name} - Cover Letter</h1>
            {"".join(f"<p>{p}</p>" for p in paragraphs)}
        </body>
        </html>
        """
        
        pdf_bytes = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.StringIO(html_content), dest=pdf_bytes)
        
        if pisa_status.err:
            logger.error(f"PDF generation error: {pisa_status.err}")
            return b""
            
        return pdf_bytes.getvalue()

    def generate_docx(self, name: str, paragraphs: list[str]) -> bytes:
        """
        Convert the cover letter paragraphs to a DOCX file using python-docx.
        """
        document = Document()
        
        # Add heading
        heading = document.add_heading(f"{name} - Cover Letter", level=1)
        
        # Add paragraphs
        for para in paragraphs:
            document.add_paragraph(para)
            
        # Save to bytes
        docx_bytes = io.BytesIO()
        document.save(docx_bytes)
        return docx_bytes.getvalue()
