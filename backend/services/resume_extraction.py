"""
Resume Extraction Service
=========================
Layered extraction pipeline with OCR fallback for any resume format.

Supports: PDF (text and scanned), DOCX, TXT, and images (PNG, JPG, JPEG).

Extraction chain (stops at first confident, non-garbage result):
  1. pdfplumber  — best for complex layouts, tables, columns
  2. pypdf       — fast fallback for simple text PDFs
  3. pdfminer.six — last resort for oddly-encoded PDFs
  4. python-docx — for .docx files
  5. Gemini Vision OCR — for scanned PDFs and raw images

A "garbage text" detector runs between each stage: if the extracted
text has too many non-printable characters or single-character words
(signature of broken glyph mappings), that stage's output is discarded
and the next strategy is tried.
"""

import io
import asyncio
import logging
import os
from typing import Optional

from schemas.models import (
    ConfidenceLevel,
    StructuredResume,
    ResumeExtractionResponse,
)

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════
# Constants
# ═══════════════════════════════════════════════════════════════

MAX_OCR_PAGES = 2
GARBAGE_THRESHOLD = 0.30  # reject if >30% garbage

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


# ═══════════════════════════════════════════════════════════════
# Garbage Text Detector
# ═══════════════════════════════════════════════════════════════

def is_garbage_text(text: str) -> bool:
    """
    Detect if extracted text is garbage (broken glyph mappings, etc.).

    Returns True if text should be treated as garbage / empty.

    Checks:
      1. Too short to be a real resume (< 50 chars)
      2. High ratio of non-printable characters (> 30 %)
      3. Too few words (< 5)
      4. High ratio of single-character "words" (> 30 %)
      5. Low ratio of alphanumeric + space characters (< 50 %)
    """
    if not text or len(text.strip()) < 50:
        return True

    stripped = text.strip()
    total = len(stripped)

    # Non-printable character ratio
    printable_count = sum(1 for c in stripped if c.isprintable() or c in "\n\r\t")
    if total > 0 and (total - printable_count) / total > GARBAGE_THRESHOLD:
        return True

    # Word-count sanity check
    words = stripped.split()
    if len(words) < 5:
        return True

    # Single-character word ratio (digits excluded — "I", "a" are ok in moderation)
    single_char_words = sum(1 for w in words if len(w) == 1 and not w.isdigit())
    if single_char_words / len(words) > GARBAGE_THRESHOLD:
        return True

    # Alphanumeric + space ratio
    alnum_count = sum(1 for c in stripped if c.isalnum() or c.isspace())
    if total > 0 and alnum_count / total < 0.5:
        return True

    return False


# ═══════════════════════════════════════════════════════════════
# Individual Extraction Strategies
# ═══════════════════════════════════════════════════════════════

def _extract_pdf_pdfplumber(pdf_bytes: bytes) -> str:
    """Strategy 1: pdfplumber (best layout-aware extraction)."""
    try:
        import pdfplumber

        text = ""
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.debug(f"pdfplumber extraction failed: {e}")
        return ""


def _extract_pdf_pypdf(pdf_bytes: bytes) -> str:
    """Strategy 2: pypdf (fast fallback for simple PDFs)."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.debug(f"pypdf extraction failed: {e}")
        return ""


def _extract_pdf_pdfminer(pdf_bytes: bytes) -> str:
    """Strategy 3: pdfminer.six (last resort for oddly-encoded PDFs)."""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract

        text = pdfminer_extract(io.BytesIO(pdf_bytes))
        return text.strip() if text else ""
    except Exception as e:
        logger.debug(f"pdfminer extraction failed: {e}")
        return ""


def _extract_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX files using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.debug(f"DOCX extraction failed: {e}")
        return ""


def _extract_plain_text(content: bytes) -> str:
    """Extract from plain text files (UTF-8 with fallback)."""
    try:
        return content.decode("utf-8", errors="replace").strip()
    except Exception as e:
        logger.debug(f"Plain text extraction failed: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════
# OCR Helpers
# ═══════════════════════════════════════════════════════════════

def _render_pdf_pages_to_images(
    pdf_bytes: bytes, max_pages: int = MAX_OCR_PAGES
) -> list[bytes]:
    """Render PDF pages to PNG byte buffers using pypdfium2."""
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(pdf_bytes)
        images: list[bytes] = []
        for i in range(min(len(pdf), max_pages)):
            page = pdf[i]
            bitmap = page.render(scale=200 / 72)  # 200 DPI
            pil_image = bitmap.to_pil()
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG")
            images.append(buf.getvalue())
        pdf.close()
        return images
    except Exception as e:
        logger.error(f"PDF → image rendering failed: {e}")
        return []


async def _ocr_with_gemini(
    images: list[bytes], llm_service
) -> str:
    """
    OCR via Gemini Vision API.

    Sends page images in parallel (asyncio.gather) and concatenates
    the transcribed text.  Capped at MAX_OCR_PAGES.
    """
    if not images or not llm_service or not llm_service.is_available:
        return ""

    if llm_service.provider != "gemini":
        logger.warning("OCR fallback requires Gemini provider — skipping")
        return ""

    try:
        import google.generativeai as genai
        from PIL import Image

        async def _ocr_page(image_bytes: bytes, page_num: int) -> str:
            try:
                img = Image.open(io.BytesIO(image_bytes))
                model = genai.GenerativeModel(model_name=llm_service.model)
                prompt = (
                    "You are an OCR engine. Transcribe ALL visible text from "
                    "this resume image faithfully and completely. Preserve the "
                    "section structure (headings, bullet points, paragraphs). "
                    "Do NOT summarize, interpret, or add any commentary. "
                    "Output ONLY the transcribed text."
                )
                response = await model.generate_content_async([prompt, img])
                return response.text if response.parts else ""
            except Exception as e:
                logger.error(f"Gemini OCR page {page_num} failed: {e}")
                return ""

        tasks = [_ocr_page(img, i) for i, img in enumerate(images)]
        results = await asyncio.gather(*tasks)
        return "\n\n".join(r for r in results if r.strip()).strip()

    except ImportError:
        logger.error("Pillow (PIL) not available — cannot run vision OCR")
        return ""
    except Exception as e:
        logger.error(f"Gemini OCR pipeline failed: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════
# Resume Extraction Service
# ═══════════════════════════════════════════════════════════════

class ResumeExtractionService:
    """
    Layered resume extraction with OCR fallback and structured parsing.

    Usage::

        from services.resume_extraction import ResumeExtractionService
        svc = ResumeExtractionService(llm_service=llm)
        result = await svc.extract(file_bytes, "resume.pdf")
        structured = await svc.extract_structured(result.text)
    """

    def __init__(self, llm_service=None):
        self.llm_service = llm_service

    # ─── Public: extract raw text ────────────────────────────

    async def extract(
        self, content: bytes, filename: str
    ) -> ResumeExtractionResponse:
        """
        Extract text from a resume file.

        Args:
            content:  Raw file bytes.
            filename: Original filename (determines format).

        Returns:
            ResumeExtractionResponse with text, confidence, warning,
            and extraction_method.
        """
        ext = os.path.splitext(filename.lower())[1]

        if ext not in SUPPORTED_EXTENSIONS:
            return ResumeExtractionResponse(
                text="",
                confidence=ConfidenceLevel.LOW,
                warning=(
                    f"Unsupported file format '{ext}'. "
                    "Please upload a PDF, DOCX, TXT, PNG, or JPG file."
                ),
                extraction_method="none",
            )

        # ── Plain text ───────────────────────────────────────
        if ext == ".txt":
            text = _extract_plain_text(content)
            if text and not is_garbage_text(text):
                return ResumeExtractionResponse(
                    text=text,
                    confidence=ConfidenceLevel.HIGH,
                    extraction_method="plain_text",
                )
            return ResumeExtractionResponse(
                text=text,
                confidence=ConfidenceLevel.LOW,
                warning=(
                    "The text file appears to be empty or contains "
                    "unreadable content."
                ),
                extraction_method="plain_text",
            )

        # ── DOCX ─────────────────────────────────────────────
        if ext == ".docx":
            text = _extract_docx(content)
            if text and not is_garbage_text(text):
                return ResumeExtractionResponse(
                    text=text,
                    confidence=ConfidenceLevel.HIGH,
                    extraction_method="python_docx",
                )
            return ResumeExtractionResponse(
                text=text,
                confidence=ConfidenceLevel.LOW,
                warning=(
                    "Could not extract readable text from this DOCX file. "
                    "Please check the file or paste your resume text manually."
                ),
                extraction_method="python_docx",
            )

        # ── Image (PNG / JPG / JPEG) ─────────────────────────
        if ext in IMAGE_EXTENSIONS:
            text = await _ocr_with_gemini([content], self.llm_service)
            if text and not is_garbage_text(text):
                return ResumeExtractionResponse(
                    text=text,
                    confidence=ConfidenceLevel.MEDIUM,
                    warning=(
                        "This resume was processed using OCR from an image. "
                        "Please review the extracted text for accuracy."
                    ),
                    extraction_method="gemini_vision_ocr",
                )
            return ResumeExtractionResponse(
                text=text,
                confidence=ConfidenceLevel.LOW,
                warning=(
                    "Could not extract readable text from this image. "
                    "Please paste your resume text manually."
                ),
                extraction_method="gemini_vision_ocr",
            )

        # ── PDF — layered text extraction ────────────────────
        strategies = [
            ("pdfplumber", _extract_pdf_pdfplumber),
            ("pypdf", _extract_pdf_pypdf),
            ("pdfminer", _extract_pdf_pdfminer),
        ]

        for method_name, strategy_fn in strategies:
            text = strategy_fn(content)
            if text and not is_garbage_text(text):
                return ResumeExtractionResponse(
                    text=text,
                    confidence=ConfidenceLevel.HIGH,
                    extraction_method=method_name,
                )
            logger.debug(
                f"PDF strategy '{method_name}' yielded garbage or "
                "empty text — trying next…"
            )

        # ── OCR fallback for scanned / image-only PDFs ───────
        logger.info(
            "All text-based PDF strategies failed. "
            "Attempting Gemini Vision OCR…"
        )
        images = _render_pdf_pages_to_images(content, max_pages=MAX_OCR_PAGES)
        if images:
            text = await _ocr_with_gemini(images, self.llm_service)
            if text and not is_garbage_text(text):
                return ResumeExtractionResponse(
                    text=text,
                    confidence=ConfidenceLevel.MEDIUM,
                    warning=(
                        "This looks like a scanned resume — we used OCR. "
                        "Please double-check the extracted text before "
                        "continuing."
                    ),
                    extraction_method="gemini_vision_ocr",
                )

        # ── All strategies exhausted ─────────────────────────
        return ResumeExtractionResponse(
            text="",
            confidence=ConfidenceLevel.LOW,
            warning=(
                "We couldn't read this file automatically. You can paste "
                "your resume text below for analysis."
            ),
            extraction_method="none",
        )

    # ─── Public: structured field extraction ─────────────────

    async def extract_structured(
        self, resume_text: str
    ) -> Optional[StructuredResume]:
        """
        Extract structured fields from resume text using regex and heuristics.
        Does not depend on the LLM.
        """
        if not resume_text or len(resume_text.strip()) < 30:
            return None

        import re
        
        # Helper to extract email
        email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", resume_text)
        email = email_match.group(0) if email_match else ""

        # Helper to extract phone
        phone_match = re.search(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", resume_text)
        phone = phone_match.group(0) if phone_match else ""
        
        # URLs
        linkedin_match = re.search(r"linkedin\.com/in/[a-zA-Z0-9_-]+", resume_text)
        linkedin = f"https://www.{linkedin_match.group(0)}" if linkedin_match else ""
        
        github_match = re.search(r"github\.com/[a-zA-Z0-9_-]+", resume_text)
        github = f"https://{github_match.group(0)}" if github_match else ""

        # Name (heuristic: first non-empty line)
        lines = [line.strip() for line in resume_text.splitlines() if line.strip()]
        name = lines[0] if lines else "Candidate"
        
        # Skills (heuristic: look for known skills)
        known_skills = ["python", "javascript", "java", "c++", "c#", "react", "node.js", "sql", "docker", "aws", "machine learning", "html", "css", "git", "linux", "agile"]
        skills = list({skill for skill in known_skills if re.search(r'\b' + re.escape(skill) + r'\b', resume_text.lower())})
        
        # Fallback empty structures for complex fields
        return StructuredResume(
            name=name[:50],
            email=email,
            phone=phone,
            linkedin=linkedin,
            github=github,
            summary="Extracted via heuristic fallback. Please refine manually.",
            skills=skills,
            experience=[],
            education=[],
            projects=[],
            certifications=[]
        )
