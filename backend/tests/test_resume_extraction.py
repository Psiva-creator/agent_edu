"""
Tests: Resume Extraction Service
=================================
Covers: clean text-layer PDF, DOCX, plain TXT, scanned/empty PDF,
corrupted file, image file, and garbage-text detector unit tests.

Every case asserts a 200-equivalent result (ExtractionResponse) with
*some* usable text or a clear low-confidence warning — never an
unhandled exception.
"""

import io
import pytest
from unittest.mock import AsyncMock, MagicMock, patch

from services.resume_extraction import (
    ResumeExtractionService,
    is_garbage_text,
    _extract_plain_text,
    _extract_docx,
    _extract_pdf_pdfplumber,
    _extract_pdf_pypdf,
)
from schemas.models import ConfidenceLevel


# ═══════════════════════════════════════════════════════════════
# Fixtures
# ═══════════════════════════════════════════════════════════════

SAMPLE_RESUME_TEXT = (
    "Jane Doe\n"
    "janedoe@email.com | San Francisco, CA\n\n"
    "SUMMARY\n"
    "Software Engineer with 4 years of experience building scalable web "
    "applications. Proficient in React, Node.js, and Python.\n\n"
    "EXPERIENCE\n"
    "Frontend Developer, Tech Corp\n"
    "2020 – Present\n"
    "- Built modern React applications improving user engagement by 30%\n"
    "- Optimized Webpack build times by 50%\n"
    "- Collaborated with UX team to redesign main dashboard\n\n"
    "EDUCATION\n"
    "B.S. Computer Science, University of California\n"
    "2014 – 2018\n\n"
    "SKILLS\n"
    "React, JavaScript, Python, Node.js, HTML, CSS, SQL\n"
)


@pytest.fixture
def extraction_service():
    """Extraction service with no LLM (tests text-based paths only)."""
    return ResumeExtractionService(llm_service=None)


@pytest.fixture
def extraction_service_with_llm():
    """Extraction service with a mocked LLM for OCR / structured tests."""
    mock_llm = MagicMock()
    mock_llm.is_available = True
    mock_llm.provider = "gemini"
    mock_llm.model = "gemini-2.5-flash-lite"
    mock_llm.generate_json = AsyncMock(return_value={
        "name": "Jane Doe",
        "email": "janedoe@email.com",
        "phone": "",
        "linkedin": "",
        "github": "",
        "summary": "Software Engineer with 4 years of experience.",
        "skills": ["React", "JavaScript", "Python", "Node.js"],
        "experience": [
            {
                "role": "Frontend Developer",
                "company": "Tech Corp",
                "dates": "2020 – Present",
                "bullets": ["Built modern React applications"],
            }
        ],
        "education": [
            {
                "degree": "B.S. Computer Science",
                "institution": "University of California",
                "year": "2018",
            }
        ],
        "projects": [],
        "certifications": [],
    })
    return ResumeExtractionService(llm_service=mock_llm)


@pytest.fixture
def sample_txt_bytes():
    """Plain text resume as bytes."""
    return SAMPLE_RESUME_TEXT.encode("utf-8")


@pytest.fixture
def sample_docx_bytes():
    """Generate a real DOCX file in memory using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("janedoe@email.com | San Francisco, CA")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "Software Engineer with 4 years of experience building "
        "scalable web applications."
    )
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("React, JavaScript, Python, Node.js, HTML, CSS, SQL")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Frontend Developer at Tech Corp (2020 – Present)")
    doc.add_paragraph(
        "Built modern React applications improving user engagement by 30%"
    )
    doc.add_heading("Education", level=2)
    doc.add_paragraph(
        "B.S. Computer Science, University of California, 2014–2018"
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_pdf_bytes():
    """Generate a real PDF with text content using pypdf."""
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DictionaryObject,
        NameObject,
        NumberObject,
        TextStringObject,
        StreamObject,
    )

    writer = PdfWriter()
    # Create a simple page with a text stream
    page = writer.add_blank_page(width=612, height=792)
    # Use reportlab if available, else add text via stream
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import letter

        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=letter)
        c.setFont("Helvetica", 12)
        y = 750
        for line in SAMPLE_RESUME_TEXT.split("\n"):
            c.drawString(72, y, line)
            y -= 16
        c.showPage()
        c.save()
        return buf.getvalue()
    except ImportError:
        # Fallback: create a minimal PDF using pypdf raw
        # (less sophisticated but still has text layer)
        content = (
            "BT /F1 12 Tf 72 750 Td "
            "(Jane Doe - Software Engineer with 4 years of experience "
            "building scalable web applications using React JavaScript "
            "Python Node.js HTML CSS SQL at Tech Corp as Frontend Developer) Tj "
            "ET"
        )
        stream = StreamObject()
        stream[NameObject("/Filter")] = NameObject("/FlateDecode")
        stream._data = content.encode("latin-1")

        resources = DictionaryObject()
        font_dict = DictionaryObject()
        font_entry = DictionaryObject()
        font_entry[NameObject("/Type")] = NameObject("/Font")
        font_entry[NameObject("/Subtype")] = NameObject("/Type1")
        font_entry[NameObject("/BaseFont")] = NameObject("/Helvetica")
        font_dict[NameObject("/F1")] = font_entry
        resources[NameObject("/Font")] = font_dict

        page[NameObject("/Resources")] = resources
        page[NameObject("/Contents")] = writer._add_object(stream)

        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()


@pytest.fixture
def empty_pdf_bytes():
    """A valid PDF with a blank page (no text layer at all)."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


@pytest.fixture
def corrupted_bytes():
    """Random bytes — not a valid file of any kind."""
    return b"\x00\x01\x02\x03GARBAGE_DATA_NOT_A_REAL_FILE\xff\xfe\xfd"


# ═══════════════════════════════════════════════════════════════
# Unit Tests: Garbage Text Detector
# ═══════════════════════════════════════════════════════════════


class TestGarbageDetector:
    """Test the is_garbage_text() function."""

    def test_none_is_garbage(self):
        assert is_garbage_text(None) is True

    def test_empty_string_is_garbage(self):
        assert is_garbage_text("") is True

    def test_short_string_is_garbage(self):
        assert is_garbage_text("too short") is True

    def test_good_resume_text_is_not_garbage(self):
        assert is_garbage_text(SAMPLE_RESUME_TEXT) is False

    def test_non_printable_heavy_text_is_garbage(self):
        # >30% non-printable characters
        garbage = "A" * 30 + "\x00" * 50 + " word " * 5
        assert is_garbage_text(garbage) is True

    def test_single_char_heavy_text_is_garbage(self):
        # >30% single-character words
        garbage = " ".join(["a"] * 20 + ["word"] * 10 + ["resume"] * 5)
        assert is_garbage_text(garbage) is True

    def test_few_words_is_garbage(self):
        assert is_garbage_text("just four words here") is True  # < 50 chars

    def test_valid_text_passes(self):
        text = (
            "John Smith has 5 years of software engineering experience "
            "at major tech companies including Google and Microsoft. "
            "He specializes in Python and JavaScript development with "
            "expertise in React, Django, and cloud platforms."
        )
        assert is_garbage_text(text) is False


# ═══════════════════════════════════════════════════════════════
# Integration Tests: Extraction Pipeline
# ═══════════════════════════════════════════════════════════════


class TestPlainTextExtraction:
    """Test .txt file extraction."""

    @pytest.mark.asyncio
    async def test_txt_extraction_succeeds(
        self, extraction_service, sample_txt_bytes
    ):
        result = await extraction_service.extract(
            sample_txt_bytes, "resume.txt"
        )
        assert result.text.strip() != ""
        assert result.confidence == ConfidenceLevel.HIGH
        assert result.extraction_method == "plain_text"
        assert result.warning is None

    @pytest.mark.asyncio
    async def test_empty_txt_returns_low_confidence(self, extraction_service):
        result = await extraction_service.extract(b"", "resume.txt")
        assert result.confidence == ConfidenceLevel.LOW
        assert result.warning is not None


class TestDocxExtraction:
    """Test .docx file extraction."""

    @pytest.mark.asyncio
    async def test_docx_extraction_succeeds(
        self, extraction_service, sample_docx_bytes
    ):
        result = await extraction_service.extract(
            sample_docx_bytes, "resume.docx"
        )
        assert result.text.strip() != ""
        assert result.confidence == ConfidenceLevel.HIGH
        assert result.extraction_method == "python_docx"
        assert "Jane Doe" in result.text or "Software" in result.text

    @pytest.mark.asyncio
    async def test_empty_docx_returns_low_confidence(self, extraction_service):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = await extraction_service.extract(buf.getvalue(), "empty.docx")
        assert result.confidence == ConfidenceLevel.LOW


class TestPdfExtraction:
    """Test PDF text extraction strategies."""

    @pytest.mark.asyncio
    async def test_text_pdf_extraction_succeeds(
        self, extraction_service, sample_pdf_bytes
    ):
        result = await extraction_service.extract(
            sample_pdf_bytes, "resume.pdf"
        )
        # Should extract something from a PDF with text
        assert result.text.strip() != ""
        assert result.confidence == ConfidenceLevel.HIGH
        assert result.extraction_method in ("pdfplumber", "pypdf", "pdfminer")

    @pytest.mark.asyncio
    async def test_empty_pdf_returns_low_confidence(
        self, extraction_service, empty_pdf_bytes
    ):
        """A blank PDF with no text should fall through all strategies."""
        result = await extraction_service.extract(
            empty_pdf_bytes, "blank.pdf"
        )
        # Without LLM, OCR cannot run — should get low confidence
        assert result.confidence == ConfidenceLevel.LOW
        assert result.warning is not None


class TestCorruptedFile:
    """Test handling of corrupted / unreadable files."""

    @pytest.mark.asyncio
    async def test_corrupted_pdf_returns_low_confidence(
        self, extraction_service, corrupted_bytes
    ):
        result = await extraction_service.extract(
            corrupted_bytes, "corrupted.pdf"
        )
        assert result.confidence == ConfidenceLevel.LOW
        assert result.warning is not None
        # Must never raise an unhandled exception
        assert isinstance(result.text, str)

    @pytest.mark.asyncio
    async def test_corrupted_docx_returns_low_confidence(
        self, extraction_service, corrupted_bytes
    ):
        result = await extraction_service.extract(
            corrupted_bytes, "corrupted.docx"
        )
        assert result.confidence == ConfidenceLevel.LOW

    @pytest.mark.asyncio
    async def test_unsupported_format_returns_warning(
        self, extraction_service
    ):
        result = await extraction_service.extract(
            b"some data", "resume.xlsx"
        )
        assert result.confidence == ConfidenceLevel.LOW
        assert "Unsupported" in result.warning


class TestImageExtraction:
    """Test image file paths (PNG/JPG) — OCR is mocked."""

    @pytest.mark.asyncio
    async def test_image_without_llm_returns_low_confidence(
        self, extraction_service
    ):
        """Without LLM service, image extraction should degrade gracefully."""
        # Create a minimal 1x1 white PNG
        from PIL import Image

        buf = io.BytesIO()
        img = Image.new("RGB", (100, 100), color="white")
        img.save(buf, format="PNG")
        result = await extraction_service.extract(
            buf.getvalue(), "resume.png"
        )
        assert result.confidence == ConfidenceLevel.LOW
        assert result.warning is not None


class TestStructuredExtraction:
    """Test structured field extraction via LLM."""

    @pytest.mark.asyncio
    async def test_structured_extraction_returns_fields(
        self, extraction_service_with_llm
    ):
        result = await extraction_service_with_llm.extract_structured(
            SAMPLE_RESUME_TEXT
        )
        assert result is not None
        assert result.name == "Jane Doe"
        assert "React" in result.skills
        assert len(result.experience) > 0
        assert len(result.education) > 0

    @pytest.mark.asyncio
    async def test_structured_extraction_without_llm_returns_none(
        self, extraction_service
    ):
        result = await extraction_service.extract_structured(
            SAMPLE_RESUME_TEXT
        )
        assert result is None

    @pytest.mark.asyncio
    async def test_structured_extraction_with_short_text_returns_none(
        self, extraction_service_with_llm
    ):
        result = await extraction_service_with_llm.extract_structured("short")
        assert result is None


# ═══════════════════════════════════════════════════════════════
# Individual Strategy Unit Tests
# ═══════════════════════════════════════════════════════════════


class TestIndividualStrategies:
    """Test individual extraction functions directly."""

    def test_plain_text_extraction(self):
        text = _extract_plain_text(b"Hello world this is a test")
        assert text == "Hello world this is a test"

    def test_plain_text_handles_encoding_errors(self):
        # Should not raise, uses errors='replace'
        text = _extract_plain_text(b"\xff\xfe\x00\x01 resume text here")
        assert isinstance(text, str)

    def test_docx_extraction_with_bad_bytes(self):
        text = _extract_docx(b"not a real docx")
        assert text == ""

    def test_pdfplumber_with_bad_bytes(self):
        text = _extract_pdf_pdfplumber(b"not a pdf")
        assert text == ""

    def test_pypdf_with_bad_bytes(self):
        text = _extract_pdf_pypdf(b"not a pdf")
        assert text == ""
